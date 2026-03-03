"""DOCX comment annotation with text-to-run mapping and fuzzy matching."""

from __future__ import annotations

import difflib
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt as docx_pt, RGBColor
from lxml import etree

from legalos.analysis.schemas import RedlineComment, RedlineOutput
from legalos.utils.errors import RedlineError

# Word XML namespaces
_NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _find_text_in_paragraphs(
    doc: Document,
    target_text: str,
    threshold: float = 0.85,
) -> tuple[int, int, int] | None:
    """Find target text in document paragraphs.

    Returns (paragraph_index, start_char, end_char) or None.
    Uses exact match first, then fuzzy fallback.
    """
    target_clean = target_text.strip()

    # Exact substring match
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text
        idx = para_text.find(target_clean)
        if idx >= 0:
            return (i, idx, idx + len(target_clean))

    # Fuzzy match fallback
    best_ratio = 0.0
    best_match = None
    for i, para in enumerate(doc.paragraphs):
        para_text = para.text
        if not para_text.strip():
            continue
        # Try matching against sliding windows of similar length
        target_len = len(target_clean)
        for start in range(0, max(1, len(para_text) - target_len + 1)):
            end = min(start + target_len + target_len // 4, len(para_text))
            window = para_text[start:end]
            # Pre-filter with quick_ratio to skip obvious non-matches
            matcher = difflib.SequenceMatcher(None, target_clean, window)
            if matcher.quick_ratio() < threshold:
                continue
            ratio = matcher.ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_match = (i, start, end)

    if best_ratio >= threshold and best_match is not None:
        return best_match
    return None


def _add_word_comment(
    doc: Document,
    para_index: int,
    comment_text: str,
    author: str,
    comment_id: int,
) -> None:
    """Add a proper Word margin comment to a paragraph using lxml.

    Injects <w:comment> into the comments part and wraps the paragraph
    content with <w:commentRangeStart>/<w:commentRangeEnd> markers.
    """
    para = doc.paragraphs[para_index]
    para_elem = para._element

    # Ensure the document has a comments part
    _ensure_comments_part(doc)

    # Add the comment definition to the comments part
    comments_part = doc.part.package.part_related_by(
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    )
    comments_elem = comments_part.element

    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    comment_elem = etree.SubElement(
        comments_elem,
        f"{{{_NSMAP['w']}}}comment",
    )
    comment_elem.set(f"{{{_NSMAP['w']}}}id", str(comment_id))
    comment_elem.set(f"{{{_NSMAP['w']}}}author", author)
    comment_elem.set(f"{{{_NSMAP['w']}}}date", now)

    # Comment body paragraph
    comment_para = etree.SubElement(comment_elem, f"{{{_NSMAP['w']}}}p")
    comment_run = etree.SubElement(comment_para, f"{{{_NSMAP['w']}}}r")
    comment_text_elem = etree.SubElement(comment_run, f"{{{_NSMAP['w']}}}t")
    comment_text_elem.text = comment_text

    # Add commentRangeStart at the beginning of the paragraph
    range_start = etree.Element(f"{{{_NSMAP['w']}}}commentRangeStart")
    range_start.set(f"{{{_NSMAP['w']}}}id", str(comment_id))
    para_elem.insert(0, range_start)

    # Add commentRangeEnd and commentReference at the end
    range_end = etree.SubElement(para_elem, f"{{{_NSMAP['w']}}}commentRangeEnd")
    range_end.set(f"{{{_NSMAP['w']}}}id", str(comment_id))

    ref_run = etree.SubElement(para_elem, f"{{{_NSMAP['w']}}}r")
    ref_elem = etree.SubElement(ref_run, f"{{{_NSMAP['w']}}}commentReference")
    ref_elem.set(f"{{{_NSMAP['w']}}}id", str(comment_id))


def _ensure_comments_part(doc: Document) -> None:
    """Ensure the DOCX package has a comments part, creating one if needed."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    # Check if comments relationship already exists
    try:
        doc.part.package.part_related_by(rel_type)
        return  # Already exists
    except KeyError:
        pass

    # Create comments part
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '</w:comments>'
    )

    comments_part = Part(
        PackURI("/word/comments.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        comments_xml.encode("utf-8"),
        doc.part.package,
    )

    doc.part.relate_to(comments_part, rel_type)


def _format_comment(comment: RedlineComment) -> str:
    """Format a redline comment for display."""
    parts = [
        f"[{comment.severity.value.upper()}]",
        f"Issue: {comment.issue}",
        f"Suggestion: {comment.suggestion}",
    ]
    if comment.alternative_language:
        parts.append(f"Alternative language: {comment.alternative_language}")
    parts.append(f"Reasoning: {comment.reasoning}")
    return " | ".join(parts)


def generate_redline(
    source_path: Path,
    redline_output: RedlineOutput,
    output_path: Path | None = None,
    author: str = "LegalOS",
) -> Path:
    """Generate an annotated DOCX with redline comments."""
    if source_path.suffix.lower() != ".docx":
        raise RedlineError(
            f"Redline requires a DOCX file, got '{source_path.suffix}'. "
            "Convert PDF to DOCX first."
        )

    try:
        doc = Document(str(source_path))
    except Exception as e:
        raise RedlineError(f"Cannot open DOCX: {e}") from e

    matched_count = 0
    unmatched_comments: list[str] = []
    comment_id = 0

    for comment in redline_output.comments:
        location = _find_text_in_paragraphs(doc, comment.target_text)
        comment_text = _format_comment(comment)

        if location is not None:
            para_idx, _, _ = location
            _add_word_comment(doc, para_idx, comment_text, author, comment_id)
            comment_id += 1
            matched_count += 1
        else:
            unmatched_comments.append(comment_text)

    # Add unmatched comments as a styled summary at the start
    if unmatched_comments:
        summary_text = f"[{author} — UNMATCHED COMMENTS]\n" + "\n\n".join(unmatched_comments)
        summary_para = doc.paragraphs[0].insert_paragraph_before(summary_text)
        # Style the summary paragraph
        for run in summary_para.runs:
            run.font.size = docx_pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
            run.font.italic = True

    if output_path is None:
        output_path = source_path.parent / f"{source_path.stem}_redlined.docx"

    doc.save(str(output_path))
    return output_path
